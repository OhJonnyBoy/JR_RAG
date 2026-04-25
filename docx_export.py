import os
import re
import uuid
import config as cfg

try:
    from docx import Document
except ImportError:
    raise ImportError("python-docx is not installed. Please run: pip install python-docx")

def export_text_to_docx(text: str, title: str = "Tailored CV Draft") -> str:
    """
    Parses basic Markdown (Headers, Bold, Bullets) and generates a Word Document.
    Returns the relative path to the generated .docx file.
    """
    doc = Document()
    doc.add_heading(title, 0)
    
    for line in text.split('\n'):
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('# '):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith('- ') or line.startswith('* '):
            # Bullet point
            p = doc.add_paragraph(style='List Bullet')
            # Handle inline bolding **text**
            _add_runs_to_paragraph(p, line[2:].strip())
        else:
            # Normal paragraph
            p = doc.add_paragraph()
            _add_runs_to_paragraph(p, line)

    export_dir = getattr(cfg, "EXPORTS_DIR", "exports")
    os.makedirs(export_dir, exist_ok=True)
    
    safe_title = "".join([c if c.isalnum() else "_" for c in title])
    filename = f"{safe_title}_{uuid.uuid4().hex[:6]}.docx"
    file_path = os.path.join(export_dir, filename)
    
    doc.save(file_path)
    return file_path

def _add_runs_to_paragraph(paragraph, text):
    """Helper to parse **bold** text and add to paragraph."""
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**') and len(part) > 4:
            paragraph.add_run(part[2:-2]).bold = True
        else:
            paragraph.add_run(part)
